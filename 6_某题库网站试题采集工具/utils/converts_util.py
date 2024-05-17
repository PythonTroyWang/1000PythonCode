import requests
from PIL import Image
from bs4 import BeautifulSoup
from io import BytesIO
import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
import uuid


def setting(doc):
    # 设置字体
    font = u'宋体'
    area = qn('w:eastAsia')
    normal = doc.styles['Normal']
    normal.font.name = font
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal._element.rPr.rFonts.set(area, font)

    # 设置页面大小
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # 设置页边距
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)


def get_last_part(string):
    pattern = r'[^/]+$'
    match = re.search(pattern, string)
    if match:
        return match.group()
    else:
        return None


def get_last_part_img(string):
    pattern = r'[^.]+$'
    match = re.search(pattern, string)
    if match:
        return match.group()
    else:
        return None


# 测试解析
def parse(content, url):
    # content = ''
    # url = 'https://v.huatu.com/tiku/analysis/170047693521222958'
    # with open('170047693521222958.html', 'r', encoding='utf-8') as file:
    #     content = file.read()

    # 创建文档对象
    document = Document()
    setting(document)

    # 一、开始解析内容
    soup = BeautifulSoup(content, 'lxml')
    title = soup.find(class_='exercise-main-title')

    # 二、获取试卷标题
    paper_title = title.text.lstrip().rstrip()
    document.add_paragraph(paper_title)
    questions = soup.find_all(class_='exercise-main-topic')

    # 三、循环获取题目所有信息
    for row in questions:
        # 1、解析题目
        jiexi_q_title(document, row)
        # 2、解析选项
        jiexi_choices(document, row)
        # 3、解析解析
        jiexi_jiexi(document, row)

    # 四、生成word文档
    output_file_name = get_last_part(url)
    document.save(f'output/{output_file_name}.docx')
    # document.save('111.docx')


def removeExif(img_bytesio):
    img_save = Image.open(img_bytesio)
    file_name = uuid.uuid4()
    suffix = img_save.format
    new_file = f'img/{file_name}.{suffix}'
    img_save.save(new_file)
    return new_file


def jiexi_q_title(document, row):
    stem = row.find(class_='main-topic-stem')
    paragraph = document.add_paragraph()
    for child in stem.children:
        if child.name == 'img':
            if child.has_attr('data-latex'):
                formula_text = child['data-latex']
                paragraph.add_run('$' + formula_text + '$')
            else:
                src = child['src']
                if 'aliyuncs.com' in src:
                    img_width = float(child['width'])
                    response = requests.get(src)
                    if response.status_code == 200:
                        img_bytesio = BytesIO(response.content)
                        img_width_cm = img_width / 25
                        if img_width_cm > 14.66:
                            print(f'普通题题干：超过边界的图片：{src}')
                            try:
                                paragraph.add_run().add_picture(img_bytesio, Cm(14.66))
                            except Exception:
                                paragraph.add_run().add_picture(removeExif(img_bytesio), Cm(14.66))
                        else:
                            print(f'普通题题干：未超过边界的图片：{src}')
                            try:
                                paragraph.add_run().add_picture(img_bytesio, Cm(img_width_cm))
                            except Exception:
                                paragraph.add_run().add_picture(removeExif(img_bytesio), Cm(img_width_cm))
                    else:
                        print(f"Failed to retrieve image from {src}")
        elif (not isinstance(child, str)) and child.find('p') is not None:
            complex_title = child.find('p')
            for complex_child in complex_title.children:
                if complex_child.name == 'img':
                    src = complex_child['src']
                    img_width = float(complex_child['width'])
                    response = requests.get(src)
                    if response.status_code == 200:
                        img_bytesio = BytesIO(response.content)
                        img_width_cm = img_width / 25
                        if img_width_cm > 14.66:
                            print(f'材料题题干：超过边界的图片：{src}')
                            try:
                                paragraph.add_run().add_picture(img_bytesio, Cm(14.66))
                            except Exception:
                                paragraph.add_run().add_picture(removeExif(img_bytesio), Cm(14.66))
                        else:
                            print(f'材料题题干：未超过边界的图片：{src}')
                            try:
                                paragraph.add_run().add_picture(img_bytesio, Cm(img_width_cm))
                            except Exception:
                                paragraph.add_run().add_picture(removeExif(img_bytesio), Cm(img_width_cm))
                    else:
                        print(f"Failed to retrieve image from {src}")
                elif isinstance(complex_child, str):
                    text = re.sub(r'\s+', '', complex_child)
                    paragraph.add_run(text)
        elif isinstance(child, str):
            text = re.sub(r'\s+', '', child)
            paragraph.add_run(text)


def jiexi_choices(document, row):
    choices = row.find(class_='main-topic-choices')
    if choices is not None:
        choice = choices.find_all(class_='main-topic-choice')
        for c in choice:
            c_text = c.text
            paragraph = document.add_paragraph(c_text)
            c_img = c.find_all('img')
            choice_img(paragraph, c_img)


def choice_img(paragraph, img):
    for i in img:
        if i.has_attr('data-latex'):
            formula_text = i['data-latex']
            paragraph.add_run('$' + formula_text + '$')
        else:
            src = i['src']
            if 'aliyuncs.com' in src:
                if src.endswith('.png') or src.endswith('jpg') or src.endswith('.jpeg'):
                    img_width = float(i['width'])
                    response = requests.get(src)
                    img_bytesio = BytesIO(response.content)
                    img_width_cm = img_width / 25
                    if img_width_cm > 14.66:
                        print(f'选项数据：超过边界的图片：{src}')
                        try:
                            paragraph.add_run().add_picture(img_bytesio, Cm(14.66))
                        except Exception:
                            paragraph.add_run().add_picture(removeExif(img_bytesio), Cm(14.66))
                    else:
                        print(f'选项数据：未超过边界的图片：{src}')
                        try:
                            paragraph.add_run().add_picture(img_bytesio, Cm(img_width_cm))
                        except Exception:
                            paragraph.add_run().add_picture(removeExif(img_bytesio), Cm(img_width_cm))


def jiexi_jiexi(document, row):
    jiexi = row.find(class_='main-topic-jiexi')
    if jiexi is not None:
        # 【答案】
        answer = jiexi.find(class_='g-right-answer-color')
        # 解析组件
        jiexi_items = jiexi.find(class_='jiexi-items')
        jiexi_item_title = jiexi_items.find_all(class_='jiexi-item-title')

        point_text = ''
        difficulty_text = ''
        region_text = ''
        source_text = ''

        for item_title in jiexi_item_title:
            # 【考点】
            if item_title.text == '考点':
                for item in list(item_title.next_siblings):
                    item = re.sub(r'\s+', '', item.string)
                    if item != '':
                        item += ';'
                        point_text += item
            # 【难度】
            if item_title.text == '难度':
                for item in list(item_title.next_siblings):
                    item = re.sub(r'\s+', '', item.string)
                    if item != '':
                        difficulty_text = item
            # 【地区】
            if item_title.text == '地区':
                for item in list(item_title.next_siblings):
                    item = re.sub(r'\s+', '', item.string)
                    if item != '':
                        region_text = item
            # 【来源】
            if item_title.text == '来源':
                for item in list(item_title.next_siblings):
                    item = re.sub(r'\s+', '', item.string)
                    if item != '':
                        source_text = item

        document.add_paragraph(f'【答案】{answer.text}')
        document.add_paragraph(f'【考点】{point_text.lstrip()}')

        paragraph = document.add_paragraph()
        for item_title in jiexi_item_title:
            # 【分析】
            if item_title.text == '解析':
                jiexi = item_title.next_sibling.next_sibling
                paragraph.add_run('【分析】')
                for child in jiexi.children:
                    if child.name == 'img':
                        if child.has_attr('data-latex'):
                            formula_text = child['data-latex']
                            paragraph.add_run('$' + formula_text + '$')
                        else:
                            src = child['src']
                            if 'aliyuncs.com' in src:
                                img_width = float(child['width'])
                                response = requests.get(src)
                                if response.status_code == 200:
                                    img_width_cm = img_width / 25
                                    img_bytesio = BytesIO(response.content)
                                    if img_width_cm > 14.66:
                                        print(f'解析数据：超过边界的图片：{src}')
                                        try:
                                            paragraph.add_run().add_picture(img_bytesio, Cm(14.66))
                                        except Exception:
                                            paragraph.add_run().add_picture(removeExif(img_bytesio), Cm(14.66))
                                    else:
                                        print(f'解析数据：未超过边界的图片：{src}')
                                        try:
                                            paragraph.add_run().add_picture(img_bytesio, Cm(img_width_cm))
                                        except Exception:
                                            paragraph.add_run().add_picture(removeExif(img_bytesio), Cm(img_width_cm))
                                else:
                                    print(f"Failed to retrieve image from {src}")
                    elif isinstance(child, str):
                        text = re.sub(r'\s+', '', child)
                        if text != '':
                            paragraph.add_run(text)

            if item_title.text == '拓展':
                tuozhan = item_title.next_sibling.next_sibling
                for child in tuozhan.children:
                    if child.name == 'img':
                        if child.has_attr('data-latex'):
                            formula_text = child['data-latex']
                            paragraph.add_run('$' + formula_text + '$')
                        else:
                            src = child['src']
                            if 'aliyuncs.com' in src:
                                img_width = float(child['width'])
                                response = requests.get(src)
                                if response.status_code == 200:
                                    img_width_cm = img_width / 25
                                    img_bytesio = BytesIO(response.content)
                                    if img_width_cm > 14.66:
                                        print(f'拓展数据：超过边界的图片：{src}')
                                        try:
                                            paragraph.add_run().add_picture(img_bytesio, Cm(14.66))
                                        except Exception:
                                            paragraph.add_run().add_picture(removeExif(img_bytesio), Cm(14.66))
                                    else:
                                        print(f'拓展数据：未超过边界的图片：{src}')
                                        try:
                                            paragraph.add_run().add_picture(img_bytesio, Cm(img_width_cm))
                                        except Exception:
                                            paragraph.add_run().add_picture(removeExif(img_bytesio), Cm(img_width_cm))
                                else:
                                    print(f"Failed to retrieve image from {src}")
                    elif isinstance(child, str):
                        text = re.sub(r'\s+', '', child)
                        if text != '':
                            paragraph.add_run(text)

        document.add_paragraph(f'【解答】{answer.text}')
        document.add_paragraph(f'【难度】{difficulty_text}')
        document.add_paragraph(f'【地区】{region_text}')
        document.add_paragraph(f'【来源】{source_text.strip()}')
