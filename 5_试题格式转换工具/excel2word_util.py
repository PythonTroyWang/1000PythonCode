import pandas as pd
from docx import Document
from docx.oxml.ns import qn
import random
import re

# 创建待生成的文档
from docx.shared import Pt, RGBColor


def setting(doc):
    # 设置字体
    font = u'宋体'
    area = qn('w:eastAsia')
    normal = doc.styles['Normal']
    normal.font.name = font
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal._element.rPr.rFonts.set(area, font)


def excel2word_with_file(file_path, q_random, s_random):
    document_q = Document()
    document_a = Document()
    setting(document_q)
    setting(document_a)
    # 读取数据
    df = pd.read_excel(file_path, engine="openpyxl")
    if q_random == "1":
        df = df.sample(n=df.index.size, replace=False, random_state=None)
        df = df.reset_index(drop=True)

    print(df)
    if s_random == "0":
        for index, row in df.iterrows():
            number = str(index + 1)
            print("正在处理第" + number + "条数据")
            if q_random == "1":
                # 处理题目
                q_row = row['题目-文字']
                temp = re.sub('^\d+', '', q_row)
                q_row_paragraph = number + temp
                document_q.add_paragraph(q_row_paragraph)
            elif q_random == "0":
                document_q.add_paragraph(row['题目-文字'])
            qa = row['选择题选项'].split('\n')
            for i in qa:
                document_q.add_paragraph(i)

            if pd.isnull(row['解析']):
                tip = "暂无解析"
                sentence = f'{number}.答案:{row["答案/正确选项"]}.{tip}'
            else:
                sentence = f'{number}.答案:{row["答案/正确选项"]}.{row["解析"]}'
            document_a.add_paragraph(sentence)
    elif s_random == "1":
        for index, row in df.iterrows():
            number = str(index + 1)
            if q_random == "1":
                # 处理题目
                q_row = row['题目-文字']
                temp = re.sub('^\d+', '', q_row)
                q_row_paragraph = number + temp
                document_q.add_paragraph(q_row_paragraph)
            elif q_random == "0":
                document_q.add_paragraph(row['题目-文字'])
            # 处理选项
            qa = row['选择题选项'].split('\n')
            answer = row["答案/正确选项"]

            data = []
            for i in qa:
                if answer in i:
                    match_str = re.match('(\w).(.*)', i)
                    data.append([match_str.group(2), 1, match_str.group(1)])
                else:
                    match_str = re.match('(\w).(.*)', i)
                    data.append([match_str.group(2), 0, match_str.group(1)])
            random.shuffle(data)

            data[0][0] = "A." + data[0][0]
            data[1][0] = "B." + data[1][0]
            data[2][0] = "C." + data[2][0]
            data[3][0] = "D." + data[3][0]

            data[0].append('A')
            data[1].append('B')
            data[2].append('C')
            data[3].append('D')

            tips = row["解析"]
            correct_answer = ''
            record_location = []
            for j in data:
                document_q.add_paragraph(j[0])
                record_location.append([j[3], findAllSubStrIndex(j[2], tips)])
                if j[1] == 1:
                    correct_answer = j[3]
            print(record_location)
            for record in record_location:
                if len(record[1]) != 0:
                    for c in record[1]:
                        tips = replace_char(tips, record[0], c)
            sentence = f'{number}.答案:{correct_answer}.{tips}'
            document_a.add_paragraph(sentence)

    # 保存文档
    document_q.save('word_result/题目.docx')
    document_a.save('word_result/答案.docx')


def replace_char(string, char, index):
    string = list(string)
    string[index] = char
    return ''.join(string)


def findAllSubStrIndex(substr, target_str):
    index_list = []
    times = target_str.count(substr)
    i = 0
    index = -1
    while i < times:
        index = target_str.find(substr, index + 1)
        index_list.append(index)
        i += 1
    return index_list
