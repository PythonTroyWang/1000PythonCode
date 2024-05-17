from docx import Document
import openpyxl
from openpyxl.styles import Alignment
import re
import os


def word2excel_with_file(file_path, num):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '执行结果'
    ws.column_dimensions['B'].width = 50
    ws.column_dimensions['G'].width = 50
    ws.column_dimensions['H'].width = 10
    ws.column_dimensions['K'].width = 100

    ws['A1'] = '类型'
    ws['B1'] = '题目-文字'
    ws['C1'] = '题目-图片链接'
    ws['D1'] = '题目-语音链接'
    ws['E1'] = '题目-视频-封面链接'
    ws['F1'] = '题目-视频-链接'
    ws['G1'] = '选择题选项'
    ws['H1'] = '答案/正确选项'
    ws['I1'] = '总分数/每空分数'
    ws['J1'] = '漏选分数'
    ws['K1'] = '解析'
    ws['L1'] = '简答题图片答案链接'

    data = []

    print("正在处理：%s" % file_path)
    doc = Document(file_path)
    pl = [paragraph.text for paragraph in doc.paragraphs]

    for i in pl:
        if len(i) == 0:
            continue
        data.append(i)

    step = 6
    b = [data[i:i + 6] for i in range(0, len(data), step)]
    k = 2
    for index, i in enumerate(b):
        p = re.compile(r'[ABCD].')
        result_match = p.match(i[0])
        if result_match is not None:
            return index + 1, '题目、选项或答案'

        ws['B' + str(k)].alignment = Alignment(horizontal='left', vertical='center', wrapText=True)
        ws['G' + str(k)].alignment = Alignment(horizontal='left', vertical='center', wrapText=True)
        ws['H' + str(k)].alignment = Alignment(horizontal='center', vertical='center', wrapText=True)
        ws['K' + str(k)].alignment = Alignment(horizontal='left', vertical='center', wrapText=True)

        p = re.compile(r'\d+.')
        result_match = p.match(i[0])
        if result_match is not None:
            ws['B' + str(k)] = i[0]
        else:
            ws['B' + str(k)] = str((int(num) + index)) + '.' + i[0]
        ws['G' + str(k)] = i[1] + '\n' + i[2] + '\n' + i[3] + '\n' + i[4]

        p = re.compile(r'答案:')
        result_match = p.match(i[5])
        if result_match is None:
            return index + 1, '题目、选项或答案'
        p = re.compile(r'答案:([ABCD])')
        result_answer = p.match(i[5]).group(1)
        ws['H' + str(k)] = result_answer
        p = re.compile(r'答案:([ABCD]).(.*)', re.S)
        result_tip = p.match(i[5]).group(2)
        ws['K' + str(k)] = result_tip
        k += 1

    wb.save('excel_result/output.xlsx')
    return 0, None


def word2excel_with2file(file_question, file_answer, num):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '执行结果'
    ws.column_dimensions['B'].width = 50
    ws.column_dimensions['G'].width = 50
    ws.column_dimensions['H'].width = 10
    ws.column_dimensions['K'].width = 100

    ws['A1'] = '类型'
    ws['B1'] = '题目-文字'
    ws['C1'] = '题目-图片链接'
    ws['D1'] = '题目-语音链接'
    ws['E1'] = '题目-视频-封面链接'
    ws['F1'] = '题目-视频-链接'
    ws['G1'] = '选择题选项'
    ws['H1'] = '答案/正确选项'
    ws['I1'] = '总分数/每空分数'
    ws['J1'] = '漏选分数'
    ws['K1'] = '解析'
    ws['L1'] = '简答题图片答案链接'

    data_question = []
    data_answer = []

    # print("正在处理：%s 和 %s" % (file_question, file_answer))
    doc_question = Document(file_question)
    pl_question = [paragraph.text for paragraph in doc_question.paragraphs]

    for i in pl_question:
        if len(i) == 0:
            continue
        data_question.append(i)

    step = 5
    b = [data_question[i:i + 5] for i in range(0, len(data_question), step)]
    k = 2
    for index, i in enumerate(b):
        p = re.compile(r'[ABCD].')
        result_match = p.match(i[0])
        if result_match is not None:
            return index, '题目或选项'

        ws['B' + str(k)].alignment = Alignment(horizontal='left', vertical='center', wrapText=True)
        ws['G' + str(k)].alignment = Alignment(horizontal='left', vertical='center', wrapText=True)

        p = re.compile(r'\d+.')
        result_match = p.match(i[0])
        if result_match is not None:
            ws['B' + str(k)] = i[0]
        else:
            ws['B' + str(k)] = str((int(num) + index)) + '.' + i[0]
        ws['G' + str(k)] = i[1] + '\n' + i[2] + '\n' + i[3] + '\n' + i[4]
        k += 1

    doc_answer = Document(file_answer)
    pl_answer = [paragraph.text for paragraph in doc_answer.paragraphs]

    for i in pl_answer:
        if len(i) == 0:
            continue
        data_answer.append(i)

    k = 2
    for index, i in enumerate(data_answer):
        p = re.compile(r'(\d+.答案)|(答案)')
        result_match = p.match(i)
        if result_match is None:
            return index, '答案'

        ws['H' + str(k)].alignment = Alignment(horizontal='center', vertical='center', wrapText=True)
        ws['K' + str(k)].alignment = Alignment(horizontal='left', vertical='center', wrapText=True)

        p = re.compile(r'\d+.答案')
        result_match = p.match(i)
        if result_match is not None:
            p = re.compile(r'\d+.答案:([ABCD])')
            result_answer = p.match(i)
            if result_answer is not None:
                p = re.compile(r'\d+.答案:([ABCD])')
                result_answer = p.match(i).group(1)
                ws['H' + str(k)] = result_answer
                p = re.compile(r'\d+.答案:([ABCD]).(.*)', re.S)
                result_tip = p.match(i).group(2)
                ws['K' + str(k)] = result_tip
            else:
                return index + 1, '答案没有选项'
        else:
            p = re.compile(r'答案:([ABCD])')
            result_answer = p.match(i)
            if result_answer is not None:
                p = re.compile(r'答案:([ABCD])')
                result_answer = p.match(i).group(1)
                ws['H' + str(k)] = result_answer
                p = re.compile(r'答案:([ABCD]).(.*)', re.S)
                result_tip = p.match(i).group(2)
                ws['K' + str(k)] = result_tip
            else:
                return index + 1, '答案没有选项'
        k += 1

    wb.save('excel_result/output.xlsx')
    return 0, None


def getFile(dir_path):
    file_type = ".docx"
    ret_array = []

    for dir_path, dir_names, file_names in os.walk(dir_path):
        for filename in file_names:
            if file_type in filename:
                ret_array.append(os.path.join(dir_path, filename))
    return ret_array
